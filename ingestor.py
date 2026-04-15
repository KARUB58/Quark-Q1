"""
╔══════════════════════════════════════════════════════════════════╗
║               QUARK Q1 — INGESTOR MODULE                        ║
║       Autonomous RAG • Dosya Yutma • Vektör Dönüşümü             ║
║     PDF / DOCX / TXT • Metin Parçalama • FileWatcher             ║
╚══════════════════════════════════════════════════════════════════╝

Ingestor modülü, yerel dosyaları okur, parçalara ayırır ve
ChromaDB vektör veritabanına kaydeder. Q1 bu sayede internete
bağlanmadan dosya içeriklerini "ezberler" ve sorulara RAG
(Retrieval-Augmented Generation) ile yanıt verir.

Bölümler:
    1. FormatConfig       — Desteklenen dosya formatları ve sabitleri
    2. TextExtractor      — PDF / DOCX / TXT metin çıkarıcı
    3. TextChunker        — Metin parçalama (chunking) motoru
    4. VectorStore        — ChromaDB vektör yazma / okuma
    5. FileWatcher        — Klasör izleme (watchdog)
    6. IngestStats        — İstatistik ve raporlama
    7. IngestManager      — Ana yönetici sınıfı

Yazar  : Quark Technologies
Sürüm  : 1.0.0
"""

import os
import sys
import time
import hashlib
import uuid
import threading
from datetime import datetime, timezone
from pathlib import Path


# ════════════════════════════════════════════════════════════════
#  ÖZ-FARKINDALIK: Modül Kontrolü
# ════════════════════════════════════════════════════════════════

INGEST_MISSING_MODULES = []

try:
    import chromadb
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    INGEST_MISSING_MODULES.append(("chromadb", "pip install chromadb"))

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    INGEST_MISSING_MODULES.append(("pypdf", "pip install pypdf"))

try:
    import docx as python_docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    INGEST_MISSING_MODULES.append(("python-docx", "pip install python-docx"))

try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    WATCHDOG_AVAILABLE = True
except ImportError:
    WATCHDOG_AVAILABLE = False
    INGEST_MISSING_MODULES.append(("watchdog", "pip install watchdog"))


def print_ingest_missing():
    """Eksik modülleri raporlar."""
    if not INGEST_MISSING_MODULES:
        return
    print("\n⚠️  [INGESTOR ÖZ-FARKINDALIK] Bazı modüller eksik:")
    for mod_name, install_cmd in INGEST_MISSING_MODULES:
        print(f"   ❌ {mod_name} → {install_cmd}")
    print()


# ════════════════════════════════════════════════════════════════
#  BÖLÜM 1 — FORMAT YAPILANDIRMASI
# ════════════════════════════════════════════════════════════════

class FormatConfig:
    """Desteklenen dosya formatları ve sabitler."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".py", ".json", ".csv"}

    # Parçalama ayarları
    CHUNK_SIZE           = 500     # kelime
    CHUNK_OVERLAP        = 50      # kelime (parçalar arası örtüşme)
    MIN_CHUNK_SIZE       = 20      # minimum kelime (çok kısa parçaları atla)

    # Vektör DB ayarları
    INGEST_DB_PATH       = "./q1_memory"
    INGEST_COLLECTION    = "q1_documents"

    # FileWatcher ayarları
    WATCH_DIR            = "./brain_data"
    WATCH_INTERVAL       = 2.0      # saniye
    WATCH_RECURSIVE      = True

    # Dosya boyut limiti
    MAX_FILE_SIZE_MB     = 50

    @staticmethod
    def is_supported(filepath: str) -> bool:
        """Dosya formatının desteklenip desteklenmediğini kontrol eder."""
        ext = os.path.splitext(filepath)[1].lower()
        return ext in FormatConfig.SUPPORTED_EXTENSIONS

    @staticmethod
    def get_file_type(filepath: str) -> str:
        """Dosya türünü döner."""
        ext = os.path.splitext(filepath)[1].lower()
        type_map = {
            ".pdf": "PDF",
            ".docx": "DOCX",
            ".txt": "TXT",
            ".md": "MARKDOWN",
            ".py": "PYTHON",
            ".json": "JSON",
            ".csv": "CSV",
        }
        return type_map.get(ext, "UNKNOWN")


# ════════════════════════════════════════════════════════════════
#  BÖLÜM 2 — METİN ÇIKARICI (Text Extractor)
# ════════════════════════════════════════════════════════════════

class TextExtractor:
    """
    Farklı dosya formatlarından düz metin çıkarır.
    Her format için ayrı okuyucu metodu.
    """

    @staticmethod
    def extract(filepath: str) -> dict:
        """
        Dosyadan metin çıkarır.

        Dönüş:
            {
                "text": str,
                "pages": int,
                "file_type": str,
                "file_name": str,
                "file_size_kb": float,
                "success": bool,
                "error": str or None,
            }
        """
        if not os.path.exists(filepath):
            return {
                "text": "", "pages": 0, "file_type": "UNKNOWN",
                "file_name": os.path.basename(filepath),
                "file_size_kb": 0, "success": False,
                "error": f"Dosya bulunamadı: {filepath}",
            }

        file_size_kb = os.path.getsize(filepath) / 1024
        file_name = os.path.basename(filepath)
        file_type = FormatConfig.get_file_type(filepath)

        # Boyut kontrolü
        if file_size_kb > FormatConfig.MAX_FILE_SIZE_MB * 1024:
            return {
                "text": "", "pages": 0, "file_type": file_type,
                "file_name": file_name, "file_size_kb": file_size_kb,
                "success": False,
                "error": f"Dosya çok büyük: {file_size_kb / 1024:.1f} MB (limit: {FormatConfig.MAX_FILE_SIZE_MB} MB)",
            }

        ext = os.path.splitext(filepath)[1].lower()

        try:
            if ext == ".pdf":
                text, pages = TextExtractor._extract_pdf(filepath)
            elif ext == ".docx":
                text, pages = TextExtractor._extract_docx(filepath)
            elif ext in (".txt", ".md", ".py", ".json", ".csv"):
                text, pages = TextExtractor._extract_text(filepath)
            else:
                return {
                    "text": "", "pages": 0, "file_type": file_type,
                    "file_name": file_name, "file_size_kb": file_size_kb,
                    "success": False,
                    "error": f"Desteklenmeyen format: {ext}",
                }

            return {
                "text": text, "pages": pages, "file_type": file_type,
                "file_name": file_name, "file_size_kb": round(file_size_kb, 2),
                "success": True, "error": None,
            }

        except Exception as e:
            return {
                "text": "", "pages": 0, "file_type": file_type,
                "file_name": file_name, "file_size_kb": file_size_kb,
                "success": False, "error": str(e),
            }

    @staticmethod
    def _extract_pdf(filepath: str) -> tuple:
        """PDF dosyasından metin çıkarır."""
        if not PYPDF_AVAILABLE:
            raise ImportError("pypdf modülü yüklü değil. Kur: pip install pypdf")

        reader = PdfReader(filepath)
        pages = len(reader.pages)
        text_parts = []

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Sayfa {i + 1}]\n{page_text.strip()}")

        full_text = "\n\n".join(text_parts)
        return full_text, pages

    @staticmethod
    def _extract_docx(filepath: str) -> tuple:
        """DOCX dosyasından metin çıkarır."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx modülü yüklü değil. Kur: pip install python-docx")

        doc = python_docx.Document(filepath)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Tabloları da oku
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        pages = max(1, len(full_text) // 3000)  # Yaklaşık sayfa tahmini
        return full_text, pages

    @staticmethod
    def _extract_text(filepath: str) -> tuple:
        """Düz metin dosyasından metin çıkarır (TXT, MD, PY, JSON, CSV)."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1254", "cp1252"]

        for enc in encodings:
            try:
                with open(filepath, "r", encoding=enc) as f:
                    text = f.read()
                pages = max(1, len(text) // 3000)
                return text.strip(), pages
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ValueError(f"Dosya kodlaması okunamadı: {filepath}")


# ════════════════════════════════════════════════════════════════
#  BÖLÜM 3 — METİN PARÇALAMA (Text Chunker)
# ════════════════════════════════════════════════════════════════

class TextChunker:
    """
    Büyük metinleri vektör veritabanı için uygun parçalara ayırır.
    Örtüşmeli (overlapping) parçalama ile bağlam kaybını önler.
    """

    @staticmethod
    def chunk_text(text: str,
                   chunk_size: int = FormatConfig.CHUNK_SIZE,
                   overlap: int = FormatConfig.CHUNK_OVERLAP,
                   min_size: int = FormatConfig.MIN_CHUNK_SIZE) -> list:
        """
        Metni kelime bazlı parçalara ayırır.

        Parametreler:
            text       : Parçalanacak metin
            chunk_size : Her parçadaki kelime sayısı
            overlap    : Parçalar arası örtüşen kelime sayısı
            min_size   : Minimum parça boyutu (altındakiler atlanır)

        Dönüş:
            [
                {"text": str, "index": int, "word_count": int, "char_count": int},
                ...
            ]
        """
        if not text or not text.strip():
            return []

        words = text.split()
        total_words = len(words)

        if total_words <= chunk_size:
            return [{
                "text": text.strip(),
                "index": 0,
                "word_count": total_words,
                "char_count": len(text.strip()),
            }]

        chunks = []
        start = 0
        chunk_index = 0

        while start < total_words:
            end = min(start + chunk_size, total_words)
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            word_count = len(chunk_words)
            if word_count >= min_size:
                chunks.append({
                    "text": chunk_text,
                    "index": chunk_index,
                    "word_count": word_count,
                    "char_count": len(chunk_text),
                })
                chunk_index += 1

            # Sonraki parçanın başlangıcı (örtüşme ile)
            start = end - overlap
            if start >= total_words:
                break
            # Sonsuz döngü önlemi
            if end == total_words:
                break

        return chunks

    @staticmethod
    def chunk_by_paragraphs(text: str, max_chunk_size: int = FormatConfig.CHUNK_SIZE) -> list:
        """
        Paragraf bazlı parçalama.
        Doğal kırılma noktalarını korur.
        """
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = []
        current_word_count = 0
        chunk_index = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_words = len(para.split())

            if current_word_count + para_words > max_chunk_size and current_chunk:
                chunk_text = "\n\n".join(current_chunk)
                chunks.append({
                    "text": chunk_text,
                    "index": chunk_index,
                    "word_count": current_word_count,
                    "char_count": len(chunk_text),
                })
                chunk_index += 1
                current_chunk = []
                current_word_count = 0

            current_chunk.append(para)
            current_word_count += para_words

        # Son parça
        if current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            if current_word_count >= FormatConfig.MIN_CHUNK_SIZE:
                chunks.append({
                    "text": chunk_text,
                    "index": chunk_index,
                    "word_count": current_word_count,
                    "char_count": len(chunk_text),
                })

        return chunks

    @staticmethod
    def get_chunk_stats(chunks: list) -> dict:
        """Parçalama istatistikleri."""
        if not chunks:
            return {"count": 0, "avg_words": 0, "total_words": 0}

        word_counts = [c["word_count"] for c in chunks]
        return {
            "count": len(chunks),
            "avg_words": round(sum(word_counts) / len(word_counts), 1),
            "total_words": sum(word_counts),
            "min_words": min(word_counts),
            "max_words": max(word_counts),
        }


# ════════════════════════════════════════════════════════════════
#  BÖLÜM 4 — VEKTÖR DEPOLAMA (Vector Store)
# ════════════════════════════════════════════════════════════════

class VectorStore:
    """
    ChromaDB üzerinde belge parçalarını vektör olarak depolar.
    RAG sorguları için benzerlik araması yapar.
    """

    def __init__(self, db_path: str = FormatConfig.INGEST_DB_PATH,
                 collection_name: str = FormatConfig.INGEST_COLLECTION):
        self.db_path = db_path
        self.collection_name = collection_name
        self.is_ready = False
        self.collection = None

        if not CHROMADB_AVAILABLE:
            return

        try:
            self.client = chromadb.PersistentClient(path=self.db_path)
            self.collection = self.client.get_or_create_collection(
                name=self.collection_name
            )
            self.is_ready = True
        except Exception as e:
            print(f"⚠️ VectorStore başlatılamadı: {e}")

    def store_chunks(self, chunks: list, file_metadata: dict) -> int:
        """
        Parçaları vektör veritabanına yazar.

        Dönüş: Başarıyla yazılan parça sayısı
        """
        if not self.is_ready or not chunks:
            return 0

        stored = 0
        file_name = file_metadata.get("file_name", "unknown")
        file_type = file_metadata.get("file_type", "UNKNOWN")
        file_hash = file_metadata.get("file_hash", "")

        for chunk in chunks:
            try:
                chunk_id = f"{file_hash}_{chunk['index']}"
                metadata = {
                    "file_name": file_name,
                    "file_type": file_type,
                    "chunk_index": chunk["index"],
                    "word_count": chunk["word_count"],
                    "ingested_at": datetime.now(timezone.utc).isoformat(),
                    "source": "ingestor",
                }

                self.collection.upsert(
                    documents=[chunk["text"]],
                    metadatas=[metadata],
                    ids=[chunk_id],
                )
                stored += 1

            except Exception:
                continue

        return stored

    def search(self, query: str, n_results: int = 5) -> list:
        """
        Sorguya en yakın belge parçalarını döner (RAG).
        """
        if not self.is_ready:
            return []

        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results,
                include=["documents", "distances", "metadatas"],
            )

            items = []
            docs = results.get("documents", [[]])[0]
            dists = results.get("distances", [[]])[0]
            metas = results.get("metadatas", [[]])[0]

            for doc, dist, meta in zip(docs, dists, metas):
                items.append({
                    "content": doc,
                    "distance": dist,
                    "relevance": max(0.0, 1.0 - dist),
                    "file_name": meta.get("file_name", "?"),
                    "file_type": meta.get("file_type", "?"),
                    "chunk_index": meta.get("chunk_index", 0),
                })

            return items

        except Exception:
            return []

    def get_document_count(self) -> int:
        """Toplam parça sayısını döner."""
        if not self.is_ready:
            return 0
        try:
            return self.collection.count()
        except Exception:
            return 0

    def get_ingested_files(self) -> list:
        """Yutulmuş dosya isimlerini döner."""
        if not self.is_ready:
            return []

        try:
            all_data = self.collection.get(include=["metadatas"])
            metas = all_data.get("metadatas", [])
            file_names = set()
            for m in metas:
                if m and "file_name" in m:
                    file_names.add(m["file_name"])
            return sorted(file_names)
        except Exception:
            return []

    def get_stats(self) -> dict:
        """Vektör deposu istatistikleri."""
        return {
            "is_ready": self.is_ready,
            "total_chunks": self.get_document_count(),
            "ingested_files": self.get_ingested_files(),
            "db_path": self.db_path,
            "collection": self.collection_name,
        }


# ════════════════════════════════════════════════════════════════
#  BÖLÜM 5 — DOSYA İZLEYİCİ (File Watcher)
# ════════════════════════════════════════════════════════════════

class _IngestEventHandler(FileSystemEventHandler if WATCHDOG_AVAILABLE else object):
    """Yeni dosya algılandığında ingest tetikler."""

    def __init__(self, ingest_callback):
        if WATCHDOG_AVAILABLE:
            super().__init__()
        self.ingest_callback = ingest_callback
        self.processed_files = set()
        self._lock = threading.Lock()

    def on_created(self, event):
        """Yeni dosya oluşturulduğunda tetiklenir."""
        if event.is_directory:
            return

        filepath = event.src_path
        if not FormatConfig.is_supported(filepath):
            return

        with self._lock:
            if filepath in self.processed_files:
                return
            self.processed_files.add(filepath)

        # Dosya yazımının tamamlanmasını bekle
        time.sleep(1.0)

        try:
            print(f"\n📂 [WATCHER] Yeni dosya algılandı: {os.path.basename(filepath)}")
            self.ingest_callback(filepath)
        except Exception as e:
            print(f"⚠️ [WATCHER] İşleme hatası: {e}")

    def on_modified(self, event):
        """Dosya değiştirildiğinde tetiklenir."""
        if event.is_directory:
            return

        filepath = event.src_path
        if not FormatConfig.is_supported(filepath):
            return

        # Değişiklikte de yeniden ingest yap
        time.sleep(1.0)
        try:
            print(f"\n📝 [WATCHER] Dosya güncellendi: {os.path.basename(filepath)}")
            self.ingest_callback(filepath)
        except Exception:
            pass


class FileWatcher:
    """
    Watchdog ile belirli bir klasörü izler.
    Yeni veya değiştirilmiş dosyaları otomatik ingest eder.
    """

    def __init__(self, watch_dir: str = FormatConfig.WATCH_DIR,
                 ingest_callback=None):
        self.watch_dir = os.path.abspath(watch_dir)
        self.ingest_callback = ingest_callback
        self.observer = None
        self.is_running = False

    def start(self) -> bool:
        """Klasör izlemeyi başlatır."""
        if not WATCHDOG_AVAILABLE:
            print("⚠️ watchdog modülü yüklü değil. Kur: pip install watchdog")
            return False

        if self.is_running:
            return True

        # Klasörü oluştur (yoksa)
        os.makedirs(self.watch_dir, exist_ok=True)

        handler = _IngestEventHandler(self.ingest_callback)
        self.observer = Observer()
        self.observer.schedule(handler, self.watch_dir, recursive=FormatConfig.WATCH_RECURSIVE)
        self.observer.daemon = True
        self.observer.start()
        self.is_running = True

        print(f"👁️ [WATCHER] Klasör izleme başlatıldı: {self.watch_dir}")
        return True

    def stop(self):
        """Klasör izlemeyi durdurur."""
        if self.observer and self.is_running:
            self.observer.stop()
            self.observer.join(timeout=5)
            self.is_running = False
            print("🛑 [WATCHER] Klasör izleme durduruldu.")

    def get_status(self) -> dict:
        """İzleme durumunu döner."""
        return {
            "is_running": self.is_running,
            "watch_dir": self.watch_dir,
            "watchdog_available": WATCHDOG_AVAILABLE,
        }


# ════════════════════════════════════════════════════════════════
#  BÖLÜM 6 — İSTATİSTİK VE RAPORLAMA
# ════════════════════════════════════════════════════════════════

class IngestStats:
    """İngest işlemlerinin istatistiklerini tutar."""

    def __init__(self):
        self.total_files_processed = 0
        self.total_chunks_created = 0
        self.total_errors = 0
        self.files_by_type = {}
        self.ingest_log = []

    def record_ingest(self, file_name: str, file_type: str,
                      chunks: int, success: bool, error: str = None):
        """Bir ingest işlemini kaydeder."""
        if success:
            self.total_files_processed += 1
            self.total_chunks_created += chunks
            self.files_by_type[file_type] = self.files_by_type.get(file_type, 0) + 1
        else:
            self.total_errors += 1

        self.ingest_log.append({
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "file_name": file_name,
            "file_type": file_type,
            "chunks": chunks,
            "success": success,
            "error": error,
        })

        if len(self.ingest_log) > 50:
            self.ingest_log = self.ingest_log[-50:]

    def get_report(self) -> dict:
        """İstatistik raporu."""
        return {
            "total_files": self.total_files_processed,
            "total_chunks": self.total_chunks_created,
            "total_errors": self.total_errors,
            "files_by_type": dict(self.files_by_type),
            "recent_log": self.ingest_log[-10:],
        }


# ════════════════════════════════════════════════════════════════
#  BÖLÜM 7 — ANA YÖNETİCİ (IngestManager)
# ════════════════════════════════════════════════════════════════

class IngestManager:
    """
    Dosya yutma işlemlerini orkestre eden ana sınıf.

    Kullanım:
        manager = IngestManager()
        result = manager.ingest_file("document.pdf")
        results = manager.ingest_directory("./documents")
    """

    def __init__(self):
        self.extractor = TextExtractor()
        self.chunker = TextChunker()
        self.vector_store = VectorStore()
        self.stats = IngestStats()
        self.watcher = FileWatcher(ingest_callback=self.ingest_file)

        # Modül uyarıları
        if INGEST_MISSING_MODULES:
            print_ingest_missing()

    def ingest_file(self, filepath: str) -> dict:
        """
        Tek bir dosyayı okur, parçalar ve vektör DB'ye yazar.

        Dönüş:
            {
                "success": bool,
                "file_name": str,
                "file_type": str,
                "chunks_created": int,
                "chunks_stored": int,
                "error": str or None,
            }
        """
        filepath = os.path.abspath(filepath)
        file_name = os.path.basename(filepath)

        # Format kontrolü
        if not FormatConfig.is_supported(filepath):
            error = f"Desteklenmeyen format: {os.path.splitext(filepath)[1]}"
            self.stats.record_ingest(file_name, "UNKNOWN", 0, False, error)
            return {"success": False, "file_name": file_name, "file_type": "UNKNOWN",
                    "chunks_created": 0, "chunks_stored": 0, "error": error}

        # Metin çıkar
        extraction = self.extractor.extract(filepath)
        if not extraction["success"]:
            self.stats.record_ingest(file_name, extraction["file_type"], 0, False, extraction["error"])
            return {"success": False, "file_name": file_name, "file_type": extraction["file_type"],
                    "chunks_created": 0, "chunks_stored": 0, "error": extraction["error"]}

        text = extraction["text"]
        if not text.strip():
            error = "Dosyadan metin çıkarılamadı (boş içerik)."
            self.stats.record_ingest(file_name, extraction["file_type"], 0, False, error)
            return {"success": False, "file_name": file_name, "file_type": extraction["file_type"],
                    "chunks_created": 0, "chunks_stored": 0, "error": error}

        # Parçala
        chunks = self.chunker.chunk_by_paragraphs(text)
        if not chunks:
            chunks = self.chunker.chunk_text(text)

        # Dosya hash'i (tekrar yutmayı önlemek için)
        file_hash = hashlib.md5(f"{file_name}_{len(text)}".encode()).hexdigest()[:16]

        # Vektör DB'ye yaz
        file_metadata = {
            "file_name": file_name,
            "file_type": extraction["file_type"],
            "file_hash": file_hash,
            "pages": extraction["pages"],
            "file_size_kb": extraction["file_size_kb"],
        }

        stored = self.vector_store.store_chunks(chunks, file_metadata)

        self.stats.record_ingest(file_name, extraction["file_type"], stored, True)

        print(f"✅ Yutuldu: {file_name} ({extraction['file_type']}) → {stored} parça")

        return {
            "success": True,
            "file_name": file_name,
            "file_type": extraction["file_type"],
            "chunks_created": len(chunks),
            "chunks_stored": stored,
            "pages": extraction["pages"],
            "file_size_kb": extraction["file_size_kb"],
            "error": None,
        }

    def ingest_directory(self, dir_path: str = None) -> dict:
        """
        Bir klasördeki tüm desteklenen dosyaları yutturur.

        Dönüş:
            {
                "total_files": int,
                "success_count": int,
                "error_count": int,
                "results": list,
            }
        """
        dir_path = dir_path or FormatConfig.WATCH_DIR
        dir_path = os.path.abspath(dir_path)

        if not os.path.isdir(dir_path):
            return {
                "total_files": 0, "success_count": 0,
                "error_count": 0, "results": [],
                "error": f"Klasör bulunamadı: {dir_path}",
            }

        results = []
        success_count = 0
        error_count = 0

        for root, dirs, files in os.walk(dir_path):
            for fname in files:
                fpath = os.path.join(root, fname)
                if FormatConfig.is_supported(fpath):
                    result = self.ingest_file(fpath)
                    results.append(result)
                    if result["success"]:
                        success_count += 1
                    else:
                        error_count += 1

        return {
            "total_files": len(results),
            "success_count": success_count,
            "error_count": error_count,
            "results": results,
        }

    def search_documents(self, query: str, n_results: int = 5) -> list:
        """RAG sorgusu — belge parçalarında arama yapar."""
        return self.vector_store.search(query, n_results)

    def start_watcher(self) -> bool:
        """FileWatcher'ı başlatır."""
        return self.watcher.start()

    def stop_watcher(self):
        """FileWatcher'ı durdurur."""
        self.watcher.stop()

    def get_full_status(self) -> dict:
        """Tam durum raporu."""
        return {
            "vector_store": self.vector_store.get_stats(),
            "watcher": self.watcher.get_status(),
            "stats": self.stats.get_report(),
            "missing_modules": INGEST_MISSING_MODULES,
        }


# ════════════════════════════════════════════════════════════════
#  BÖLÜM 8 — MODÜL SEVİYESİ ERİŞİM
# ════════════════════════════════════════════════════════════════

_manager = None

def get_manager() -> IngestManager:
    """IngestManager singleton erişimi."""
    global _manager
    if _manager is None:
        _manager = IngestManager()
    return _manager

def ingest_file(filepath: str) -> dict:
    """Modül seviyesinde dosya yutma."""
    return get_manager().ingest_file(filepath)

def search(query: str, n_results: int = 5) -> list:
    """Modül seviyesinde RAG araması."""
    return get_manager().search_documents(query, n_results)
